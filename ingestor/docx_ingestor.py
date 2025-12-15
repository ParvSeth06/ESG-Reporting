# docx_ingestor.py
from docx import Document
from pathlib import Path

class DOCXIngestor:
    """
    Extracts all paragraphs + tables from .docx files.
    Returns:
        {
            "text": full_text,
            "paragraphs": [...],
            "tables": [[row1, row2, ...], ...],
            "source_type": "docx"
        }
    """

    @staticmethod
    def parse(file_path: str):
        file_path = Path(file_path)
        if not file_path.exists():
            return {"error": "DOCX file not found"}

        try:
            doc = Document(file_path)

            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables = []

            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                tables.append(rows)

            return {
                "text": "\n".join(paragraphs),
                "paragraphs": paragraphs,
                "tables": tables,
                "source_type": "docx"
            }

        except Exception as e:
            return {"error": f"DOCX parse error: {str(e)}"}
