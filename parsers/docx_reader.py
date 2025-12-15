
# -----------------------------
# File: docx_reader.py
# -----------------------------
"""Starter DOCX reader utilities.
Requires: python-docx (pip install python-docx)
Functions:
- extract_docx_text(path) -> str
- extract_docx_tables(path) -> list[pandas.DataFrame]
"""

try:
    import docx
except Exception:
    docx = None


def extract_docx_text(path: str) -> str:
    if not docx:
        raise ImportError("python-docx not installed. Install with `pip install python-docx`")

    doc = docx.Document(path)
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Keep paragraph breaks
    return "\n\n".join(paras)


def extract_docx_tables(path: str) -> List[pd.DataFrame]:
    if not docx:
        raise ImportError("python-docx not installed. Install with `pip install python-docx`")

    doc = docx.Document(path)
    tables = []
    for table in doc.tables:
        data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            data.append(cells)
        df = pd.DataFrame(data)
        df = df.dropna(axis=0, how="all").dropna(axis=1, how="all")
        if df.shape[0] and df.shape[1]:
            tables.append(df)

    return tables

